import re
from docx import Document
from models.analysis_result import AnalysisResult


class DocumentAnalyzer:

    def analyze(self, filepath):

        document = Document(filepath)

        result = AnalysisResult()

        result.filename = filepath.split("\\")[-1]

        result.paragraphs = len(document.paragraphs)

        blank = 0
        word_headings = 0
        day_headings = 0
        smart_topics = 0

        for paragraph in document.paragraphs:

            text = paragraph.text.strip()

            if text == "":
                blank += 1
                continue

            # Normal Word headings
            if paragraph.style.name.startswith("Heading"):
                word_headings += 1

            # Detect DAY 1, DAY 2*, etc.
            if "DAY" in text.upper():
                print(text)

            # Detect ALL CAPS topics
            if (
                text.isupper()
                and len(text) > 8
                and not text.startswith("DAY")
                and "PRAYER POINT" not in text
            ):
                smart_topics += 1

        result.blank_paragraphs = blank
        result.word_headings = word_headings
        result.day_headings = day_headings
        result.smart_topics = smart_topics

        return result